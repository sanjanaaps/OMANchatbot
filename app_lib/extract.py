import os
import logging
from typing import Optional

# Document processing imports
try:
    import PyPDF2
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from various file formats
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text content
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        raise

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    if not PDF_AVAILABLE:
        raise ImportError("PDF processing libraries not available. Install PyPDF2 and pdfminer.six")
    
    text = ""
    
    # Try pdfminer first (better for complex PDFs)
    try:
        text = pdfminer_extract(file_path)
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"pdfminer failed for {file_path}: {str(e)}")
    
    # Fallback to PyPDF2
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        if not text.strip():
            raise ValueError("No text could be extracted from PDF")
        
        return text
    
    except Exception as e:
        logger.error(f"PyPDF2 extraction failed for {file_path}: {str(e)}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx library not available")
    
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            raise ValueError("No text could be extracted from DOCX")
        
        return text
    
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
        raise

def extract_text_from_image(file_path: str) -> str:
    """Extract text from image using OCR"""
    if not OCR_AVAILABLE:
        raise ImportError("OCR libraries not available. Install pytesseract and Pillow")
    
    try:
        # Open image
        image = Image.open(file_path)
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Extract text using Tesseract
        text = pytesseract.image_to_string(image)
        
        if not text.strip():
            raise ValueError("No text could be extracted from image")
        
        return text
    
    except Exception as e:
        logger.error(f"OCR extraction failed for {file_path}: {str(e)}")
        raise

def chunk_text(text: str, chunk_size: int = 4000, overlap: int = 200) -> list:
    """
    Split text into chunks for processing
    
    Args:
        text: Text to chunk
        chunk_size: Maximum size of each chunk
        overlap: Number of characters to overlap between chunks
        
    Returns:
        List of text chunks
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at sentence boundary
        if end < len(text):
            # Look for sentence endings within the last 200 characters
            search_start = max(start + chunk_size - 200, start)
            sentence_endings = ['.', '!', '?', '\n\n']
            
            for i in range(end - 1, search_start - 1, -1):
                if text[i] in sentence_endings:
                    end = i + 1
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks

def clean_text(text: str) -> str:
    """
    Clean extracted text by removing extra whitespace and normalizing
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    import re
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', '', text)
    
    # Remove multiple consecutive punctuation
    text = re.sub(r'([\.\,\!\?\;\:])\1+', r'\1', text)
    
    return text.strip()

def get_file_info(file_path: str) -> dict:
    """
    Get basic information about a file
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dictionary with file information
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    stat = os.stat(file_path)
    
    return {
        'filename': os.path.basename(file_path),
        'size': stat.st_size,
        'extension': os.path.splitext(file_path)[1].lower(),
        'modified': stat.st_mtime
    }